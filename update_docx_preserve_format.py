from docx import Document
from docx.oxml import OxmlElement
import os

def update_text_preserve_formatting(file_path):
    """
    Update text content ONLY while preserving ALL formatting properties
    (colors, fonts, bold, italic, sizes, etc.)
    """
    try:
        doc = Document(file_path)

        # Text replacements - these will find and replace while keeping formatting
        replacements = {
            # Title updates
            "JAVA FULL STACK DEVELOPER": "Senior Java Full Stack Engineer | Microservices | Cloud | Platform Reliability",
            "Java Full Stack Developer": "Senior Java Full Stack Engineer | Microservices | Cloud | Platform Reliability",

            # Contact info
            "603.858.7546": "603-858-7546",
            "https://www.linkedin.com/in/mohankotni/": "https://www.linkedin.com/in/ramkotni/",

            # Years of experience
            "16 years of IT experience": "18 years of experience",
            "17 years of proven experience": "18 years of experience",
            "An accomplished Java Full Stack Developer with 16 years": "Senior Java Full Stack Engineer with 18 years",
            "An accomplished Java Full Stack Developer with 17 years": "Senior Java Full Stack Engineer with 18 years",
            "A highly skilled Java Full Stack Developer with 16 years": "Senior Java Full Stack Engineer with 18 years",
            "A highly skilled Java Full Stack Developer with 17 years": "Senior Java Full Stack Engineer with 18 years",
        }

        # Process all paragraphs
        for para in doc.paragraphs:
            para_text = para.text

            # Check if this paragraph contains any text to replace
            for old_text, new_text in replacements.items():
                if old_text in para_text:
                    # Replace in runs while preserving formatting
                    for run in para.runs:
                        if old_text in run.text:
                            # Preserve formatting properties
                            run.text = run.text.replace(old_text, new_text)

        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para_text = para.text
                        for old_text, new_text in replacements.items():
                            if old_text in para_text:
                                for run in para.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)

        doc.save(file_path)
        print(f"✓ Updated (formatting preserved): {os.path.basename(file_path)}")
        return True
    except Exception as e:
        print(f"✗ Error: {os.path.basename(file_path)} - {str(e)}")
        return False

# Files to update
resume_files = [
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram-Kotni-AgilePLM.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram-Kotni-AgilePLM.docx"
]

print("Updating resume files with formatting preservation...\n")
updated = 0

for file_path in resume_files:
    if os.path.exists(file_path):
        if update_text_preserve_formatting(file_path):
            updated += 1

print(f"\n✓ Successfully updated {updated} files!")
print("✓ All original formatting preserved (colors, fonts, styles, sizes)")

