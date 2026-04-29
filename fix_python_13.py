from docx import Document

file = r'C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx'
doc = Document(file)

# Fix para 13 - the "13" is split across runs
para = doc.paragraphs[13]

print(f"Before: {para.text[:80]}...")
print(f"Runs: {len(para.runs)}")
for i, run in enumerate(para.runs):
    print(f"  Run {i}: {repr(run.text[:40])}")

if len(para.runs) >= 2:
    # The "13" is in run 1, change it to "18"
    para.runs[1].text = '18'
    print("\nAfter fix:")
    print(f"After: {para.text[:80]}...")

doc.save(file)
print("\nSaved successfully!")

# Also update the Resume folder copy
file2 = r'C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx'
doc2 = Document(file2)
para2 = doc2.paragraphs[13]
if len(para2.runs) >= 2:
    para2.runs[1].text = '18'
doc2.save(file2)
print("Saved Resume folder copy!")

