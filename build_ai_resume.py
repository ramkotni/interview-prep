"""
Build Ram Kotni — AI Engineer resume from existing Full Stack resume.
Changes:
  1. Title: "Senior Full Stack Engineer" -> "AI Engineer"
  2. Objective: updated for AI Engineer role
  3. Professional Summary: add AI strengths
  4. Technical Skills: add AI/ML row
  5. ERCOT section: add AI project (RIOO-2893 + Payment Scheduler)
  6. Key Achievements: add AI achievement
"""

import copy
from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

SRC = r"C:\RamKotni\Personal\interview-prep\FINAL-SUBMISSION-2026\Ram-Kotni-Senior-FullStack-Engineer.docx"
DST = r"C:\RamKotni\Personal\interview-prep\FINAL-SUBMISSION-2026\Ram-Kotni-AI-Engineer.docx"

SIZE_10 = 127000   # 10pt in EMU

# ─── helpers ────────────────────────────────────────────────────────────────

def set_run(run, text, bold=None, italic=None, size=SIZE_10, color_hex=None):
    run.text = text
    run.bold = bold
    run.italic = italic
    run.font.size = Emu(size)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)

def clone_paragraph_after(ref_para, new_style_name=None):
    """Insert a new blank paragraph immediately after ref_para, cloning its XML format."""
    new_p = OxmlElement('w:p')
    # Copy paragraph properties from the reference paragraph
    ref_pPr = ref_para._p.find(qn('w:pPr'))
    if ref_pPr is not None:
        new_pPr = copy.deepcopy(ref_pPr)
        new_p.insert(0, new_pPr)
    ref_para._p.addnext(new_p)
    # Wrap in Paragraph and set style
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, ref_para._p.getparent())
    if new_style_name:
        try:
            para.style = new_style_name
        except Exception:
            pass
    return para

def add_run_to_para(para, text, bold=None, italic=None, size=SIZE_10, color_hex=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Emu(size)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return run

def find_para_containing(doc, text_fragment):
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return i, p
    return None, None

# ─── load doc ────────────────────────────────────────────────────────────────

doc = Document(SRC)
paras = doc.paragraphs

# ─── 1. Change title ─────────────────────────────────────────────────────────
# P1 = "Senior Full Stack Engineer"
title_para = paras[1]
for run in title_para.runs:
    if "Senior Full Stack Engineer" in run.text:
        run.text = run.text.replace("Senior Full Stack Engineer", "AI Engineer")
        break

# ─── 2. Update Objective ─────────────────────────────────────────────────────
# P6 — objective paragraph
obj_idx, obj_para = find_para_containing(doc, "To obtain a challenging")
if obj_para:
    for run in obj_para.runs:
        run.text = ""
    # Re-build runs matching original bold pattern
    r0 = obj_para.add_run("")
    r0.bold = True; r0.font.size = Emu(SIZE_10); r0.font.color.rgb = RGBColor(0,0,0)
    r1 = obj_para.add_run("To obtain a challenging ")
    r1.bold = None; r1.font.size = Emu(SIZE_10); r1.font.color.rgb = RGBColor(0,0,0)
    r2 = obj_para.add_run("AI Engineer | Full Stack | Agentic AI | Cloud")
    r2.bold = True; r2.font.size = Emu(SIZE_10); r2.font.color.rgb = RGBColor(0,0,0)
    r3 = obj_para.add_run(
        " position in a dynamic and innovative organization, leveraging my "
    )
    r3.bold = None; r3.font.size = Emu(SIZE_10); r3.font.color.rgb = RGBColor(0,0,0)
    r4 = obj_para.add_run("18 years")
    r4.bold = True; r4.font.size = Emu(SIZE_10); r4.font.color.rgb = RGBColor(0,0,0)
    r5 = obj_para.add_run(
        " of experience building enterprise systems with Java, Angular, and Python, "
        "now applying "
    )
    r5.bold = None; r5.font.size = Emu(SIZE_10); r5.font.color.rgb = RGBColor(0,0,0)
    r6 = obj_para.add_run("Agentic AI, LLMs, GitHub Copilot, APScheduler, and pandas")
    r6.bold = True; r6.font.size = Emu(SIZE_10); r6.font.color.rgb = RGBColor(0,0,0)
    r7 = obj_para.add_run(
        " to automate complex workflows, detect anomalies, and deliver intelligent "
        "full-stack solutions at scale."
    )
    r7.bold = None; r7.font.size = Emu(SIZE_10); r7.font.color.rgb = RGBColor(0,0,0)

# ─── 3. Update Professional Summary ──────────────────────────────────────────
# P10 — first summary bullet
sum_idx, sum_para = find_para_containing(doc, "18 years of experience designing")
if sum_para:
    for run in sum_para.runs:
        run.text = ""
    r0 = sum_para.add_run("18 years")
    r0.bold = True; r0.font.size = Emu(SIZE_10); r0.font.color.rgb = RGBColor(0,0,0)
    r1 = sum_para.add_run(
        " of experience designing, developing, and automating complex enterprise systems "
        "using Java, JAX-RS/Spring Boot, Angular, Python, and AI/ML technologies. "
        "Recently pivoted to AI Engineering — building agentic workflows, payment failure "
        "detection pipelines, and LLM-assisted development using "
    )
    r1.bold = None; r1.font.size = Emu(SIZE_10); r1.font.color.rgb = RGBColor(0,0,0)
    r2 = sum_para.add_run("GitHub Copilot, APScheduler, pandas, and Oracle AI.")
    r2.bold = True; r2.font.size = Emu(SIZE_10); r2.font.color.rgb = RGBColor(0,0,0)

# Update core strengths lines
_, cs_para = find_para_containing(doc, "Core strengths")
if cs_para:
    for run in cs_para.runs:
        if "Core strengths" in run.text:
            run.text = "Core strengths"
        elif ":" in run.text:
            run.text = ": "

_, be_para = find_para_containing(doc, "Backend (70%)")
if be_para:
    for run in be_para.runs:
        run.text = ""
    r0 = be_para.add_run("Backend (60%):")
    r0.bold = True; r0.font.size = Emu(SIZE_10); r0.font.color.rgb = RGBColor(0,0,0)
    r1 = be_para.add_run(" Java 17+, JAX-RS, Spring Boot, Microservices, RESTful API design, JPA/Hibernate, secure coding.")
    r1.bold = None; r1.font.size = Emu(SIZE_10); r1.font.color.rgb = RGBColor(0,0,0)

_, fe_para = find_para_containing(doc, "Frontend (30%)")
if fe_para:
    for run in fe_para.runs:
        run.text = ""
    r0 = fe_para.add_run("AI & Automation (25%):")
    r0.bold = True; r0.font.size = Emu(SIZE_10); r0.font.color.rgb = RGBColor(0,0,0)
    r1 = fe_para.add_run(" GitHub Copilot, Agentic AI, APScheduler, pandas, LangChain, OpenAI API, Plan-driven development.")
    r1.bold = None; r1.font.size = Emu(SIZE_10); r1.font.color.rgb = RGBColor(0,0,0)

_, cd_para = find_para_containing(doc, "Cloud & DevOps:")
if cd_para:
    pass  # keep as-is

# Insert "Frontend (15%)" after Cloud & DevOps line
_, delivery_para = find_para_containing(doc, "Delivery: Agile/Scrum")
if delivery_para:
    for run in delivery_para.runs:
        run.text = ""
    r0 = delivery_para.add_run("Frontend (15%):")
    r0.bold = True; r0.font.size = Emu(SIZE_10); r0.font.color.rgb = RGBColor(0,0,0)
    r1 = delivery_para.add_run(" Angular 16+, TypeScript, RxJS, Bootstrap, responsive UI/UX.")
    r1.bold = None; r1.font.size = Emu(SIZE_10); r1.font.color.rgb = RGBColor(0,0,0)

_, dm_para = find_para_containing(doc, "Data & Messaging:")
if dm_para:
    for run in dm_para.runs:
        run.text = ""
    r0 = dm_para.add_run("Data & Delivery:")
    r0.bold = True; r0.font.size = Emu(SIZE_10); r0.font.color.rgb = RGBColor(0,0,0)
    r1 = dm_para.add_run(" Oracle, PostgreSQL, MongoDB; Kafka/RabbitMQ; Agile/Scrum, cross-functional collaboration.")
    r1.bold = None; r1.font.size = Emu(SIZE_10); r1.font.color.rgb = RGBColor(0,0,0)

# ─── 4. Update Technical Skills — add AI/ML row after Languages ──────────────
# Find Languages list paragraph and insert AI/ML after
lang_idx, lang_para = find_para_containing(doc, "Languages:")
if lang_para:
    # Update Languages to add Python (already there), keep as-is but add AI note
    for run in lang_para.runs:
        if "Java (8" in run.text or "Java (817)" in run.text:
            run.text = run.text.replace("Java (817)", "Java (8–17)").replace("Java (8", "Java (8")
    # Insert new AI/ML Tools row after Languages
    new_ai_para = clone_paragraph_after(lang_para, 'List Paragraph')
    r0 = new_ai_para.add_run("AI / ML Tools:")
    r0.bold = True; r0.font.size = Emu(SIZE_10)
    r1 = new_ai_para.add_run(" GitHub Copilot, Agentic AI Workflows, scikit-learn, Flask, joblib, APScheduler, pandas, NumPy, cx_Oracle ")
    r1.bold = None; r1.font.size = Emu(SIZE_10)

# ─── 5. Update ERCOT job title ────────────────────────────────────────────────
_, ercot_title_para = find_para_containing(doc, "Sr. Senior Java Full Stack Engineer")
if ercot_title_para:
    for run in ercot_title_para.runs:
        run.text = run.text.replace(
            "Sr. Senior Java Full Stack Engineer | Microservices | Cloud | Platform Reliability",
            "Sr. AI Engineer | Full Stack | Agentic AI | Microservices | Cloud"
        ).replace(
            "Sr. Senior Java Full Stack Engineer | Microservices | Cloud | Platform Reliabili",
            "Sr. AI Engineer | Full Stack | Agentic AI | Microservices | Cloud | "
        )

# ─── 6. Update ERCOT project line ────────────────────────────────────────────
_, ercot_proj_para = find_para_containing(doc, "RIOO-RS, GINR & Repower System")
if ercot_proj_para:
    for run in ercot_proj_para.runs:
        if "RIOO-RS, GINR & Repower System" in run.text:
            run.text = "RIOO-RS, GINR & Repower System — Agentic AI, Retirement CR Automation & Payment Monitoring"

# ─── 7. Insert AI project bullets after last ERCOT bullet ────────────────────
# Last ERCOT bullet is P51: "Collaborated in design reviews..."
_, last_ercot_bullet = find_para_containing(doc, "Collaborated in design reviews and defect triage")
if last_ercot_bullet:
    # ── Sub-header: AI Engineering Use Case 1 ───────────────────────────────
    p_sub = clone_paragraph_after(last_ercot_bullet, 'paragraph')
    r0 = p_sub.add_run("AI Engineering — Use Case 1: GINR Feasibility Screening Model (Predictive ML)")
    r0.bold = True; r0.italic = True; r0.font.size = Emu(SIZE_10); r0.font.color.rgb = RGBColor(0,0,0)

    # ── Bullet 1 ──────────────────────────────────────────────────────────────
    p1 = clone_paragraph_after(p_sub, 'paragraph')
    add_run_to_para(p1,
        "Designed and built an end-to-end GINR feasibility screening ML system — engineering "
        "20 features from 16 Oracle DB tables (ginr_request, POIB_DATA, CONNECTIVITY_NODES, "
        "MOST_LIMITING_SERIES_ELEMENTS, CONGESTION_MANAGEMENT_ZONES), training a scikit-learn "
        "Logistic Regression pipeline (ROC-AUC > 0.85, 5-fold CV), and serving real-time "
        "PASS/FAIL predictions via a Flask REST microservice (port 5050) integrated with the "
        "existing Spring Boot RARF application.",
        bold=None, color_hex="000000")

    # ── Bullet 2 ──────────────────────────────────────────────────────────────
    p2 = clone_paragraph_after(p1, 'paragraph')
    add_run_to_para(p2,
        "Architected a full ML lifecycle system: ColumnTransformer feature pipeline "
        "(StandardScaler, OneHotEncoder, SimpleImputer), versioned model registry (semantic "
        "versioning with joblib), prediction audit log, drift detection (15% FAIL-rate threshold), "
        "and automated retraining feedback loop triggered by Oracle CR status transitions "
        "(SENT_BACK / READY_FOR_MODELING). Deployed 4-tier risk scoring output "
        "(LOW/MEDIUM/HIGH/CRITICAL) with top risk factor explanations surfaced to developers at submission time.",
        bold=None, color_hex="000000")

    # ── Bullet 3 ──────────────────────────────────────────────────────────────
    p3 = clone_paragraph_after(p2, 'paragraph')
    add_run_to_para(p3,
        "Engineered derived features from Oracle data: crez_flag (PUCT CREZ county mapping), "
        "months_to_cod (COD-submitDate), project_exceeds_thermal_flag (project_mw > MIN MLSE rating), "
        "nameplate_to_maxgen_ratio (overbuild detection); implemented Flask microservice "
        "(ginr_prediction_api.py) exposing 6 endpoints: POST /api/v1/score, batch, GET score/{id}, "
        "health, model/info, drift — consumed by Spring Boot RestTemplate in ChangeRequestService.",
        bold=None, color_hex="000000")

    # ── Sub-header: Use Case 2 ────────────────────────────────────────────────
    p_sub2 = clone_paragraph_after(p3, 'paragraph')
    r0 = p_sub2.add_run("AI Engineering — Use Case 2: GINR Payment Failure Monitor (Daily Scheduler)")
    r0.bold = True; r0.italic = True; r0.font.size = Emu(SIZE_10); r0.font.color.rgb = RGBColor(0,0,0)

    # ── Bullet 4 ──────────────────────────────────────────────────────────────
    p4 = clone_paragraph_after(p_sub2, 'paragraph')
    add_run_to_para(p4,
        "Designed and built an end-to-end automated GINR payment failure monitoring pipeline "
        "in Python using APScheduler, pandas, and Oracle DB — detecting 7 categories of payment "
        "failures (unpaid FIS/SS/Analysis fees, Authorize.net sync errors, amount mismatches, "
        "duplicate payments, stale PENDING_REVIEW CRs) across 100+ daily records and delivering "
        "severity-prioritized HTML/CSV reports to operations managers.",
        bold=None, color_hex="000000")

    # ── Bullet 5 ──────────────────────────────────────────────────────────────
    p5 = clone_paragraph_after(p4, 'paragraph')
    add_run_to_para(p5,
        "Architected a production-grade BlockingScheduler with CronTrigger (7:00 AM CST daily), "
        "1-hour misfire grace window, job coalescing, max_instances=1, and SIGTERM-based graceful "
        "shutdown for Docker/systemd environments; built audit trail CSV for trend analysis "
        "and implemented --dry-run / --now CLI flags for safe QA testing.",
        bold=None, color_hex="000000")

    # ── Bullet 6 ──────────────────────────────────────────────────────────────
    p6 = clone_paragraph_after(p5, 'paragraph')
    add_run_to_para(p6,
        "Developed rich HTML email notification system using Python smtplib/MIME — dynamic "
        "severity banners (CRITICAL/HIGH/MEDIUM/ALL CLEAR), color-coded failure sections with "
        "detail tables, embedded Oracle verification SQL, and attached CSV reports delivered "
        "via ERCOT's Office 365 SMTP relay with auto-escalating subject lines.",
        bold=None, color_hex="000000")

    # ── Tech stack line ───────────────────────────────────────────────────────
    p_tech = clone_paragraph_after(p6, 'paragraph')
    r0 = p_tech.add_run("Tech: ")
    r0.bold = True; r0.font.size = Emu(SIZE_10); r0.font.color.rgb = RGBColor(0,0,0)
    r1 = p_tech.add_run(
        "Python 3.10, scikit-learn (Logistic Regression, Pipeline, ColumnTransformer), "
        "Flask, pandas, NumPy, joblib, Oracle SQL (16-table JOIN), Java 17, JAX-RS/Jersey, "
        "Spring Boot RestTemplate, JPA/Hibernate, Angular 16, TypeScript, RxJS, "
        "APScheduler, smtplib/MIME, cx_Oracle, Docker, Jenkins, Maven, GitHub Copilot (Agentic AI)"
    )
    r1.bold = None; r1.font.size = Emu(SIZE_10); r1.font.color.rgb = RGBColor(0,0,0)

# ─── 8. Update Key Achievements — add AI achievement ─────────────────────────
_, ach1_para = find_para_containing(doc, "Improved system reliability")
if ach1_para:
    p_ai_ach = clone_paragraph_after(ach1_para, 'List Paragraph')
    r0 = p_ai_ach.add_run("ML-driven feasibility screening:")
    r0.bold = True; r0.font.size = Emu(SIZE_10)
    r1 = p_ai_ach.add_run(
        " Built GINR Feasibility Screening ML system (scikit-learn Logistic Regression, ROC-AUC > 0.85) "
        "integrating 20 features from 16 Oracle DB tables — enabling real-time PASS/FAIL scoring via "
        "Flask REST API at submission time, replacing manual planner pre-screening."
    )
    r1.bold = None; r1.font.size = Emu(SIZE_10)

    p_cop_ach = clone_paragraph_after(p_ai_ach, 'List Paragraph')
    r0 = p_cop_ach.add_run("AI-driven payment automation:")
    r0.bold = True; r0.font.size = Emu(SIZE_10)
    r1 = p_cop_ach.add_run(
        " Automated 7-category GINR payment failure detection using Python/pandas APScheduler, "
        "scanning 100+ records daily and replacing manual review — delivering actionable "
        "severity-prioritized reports to managers within minutes of the 7 AM CST run."
    )
    r1.bold = None; r1.font.size = Emu(SIZE_10)

# ─── 9. Save ─────────────────────────────────────────────────────────────────
doc.save(DST)
print(f"Saved: {DST}")




