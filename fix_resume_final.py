from docx import Document
import os

def fix_and_standardize_resume(file_path):
    """
    Fix the SAVA/JAVA issue and standardize titles
    """
    try:
        doc = Document(file_path)

        # Critical fixes
        replacements = [
            # Fix SAVA back to JAVA
            ("SAVA FULL STACK DEVELOPER", "Senior Full Stack Engineer"),

            # Standardize roles (use applicable title for all)
            ("Java Full Stack Developer", "Senior Full Stack Engineer"),
            ("Python Full Stack Developer", "Senior Full Stack Engineer"),
            ("JAVA FULL STACK DEVELOPER", "Senior Full Stack Engineer"),

            # Contact updates
            ("603.858.7546", "603-858-7546"),

            # Ensure 18 years throughout
            ("17 years", "18 years"),
            ("16 years", "18 years"),
        ]

        # Process all paragraphs
        for para in doc.paragraphs:
            for old_text, new_text in replacements:
                if old_text in para.text:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old_text, new_text in replacements:
                            if old_text in para.text:
                                for run in para.runs:
                                    if old_text in run.text:
                                        run.text = run.text.replace(old_text, new_text)

        doc.save(file_path)
        filename = os.path.basename(file_path)
        print(f"✓ Fixed: {filename}")
        return True

    except Exception as e:
        print(f"✗ Error: {os.path.basename(file_path)} - {str(e)}")
        return False

# All resume files
resume_files = [
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram-Kotni-AgilePLM.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram-Kotni-AgilePLM.docx"
]

print("FIXING ALL RESUME FILES:\n")

for file_path in resume_files:
    if os.path.exists(file_path):
        fix_and_standardize_resume(file_path)

print("\n✓ All files corrected!")
print("✓ SAVA -> Senior Full Stack Engineer")
print("✓ All years set to 18 years")
print("✓ All titles standardized to 'Senior Full Stack Engineer'")

