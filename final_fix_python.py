from docx import Document

# Both files
files = [
    r'C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx',
    r'C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx'
]

for file in files:
    doc = Document(file)
    para = doc.paragraphs[6]

    # Get current text and fix all instances of 13
    full_text = para.text
    fixed_text = full_text.replace('13+', '18+')

    # Clear and rebuild
    for run in para.runs:
        run._element.getparent().remove(run._element)

    para.add_run(fixed_text)
    doc.save(file)
    print(f"Fixed: {file.split(chr(92))[-1]}")
    print(f"  New text: {para.text[50:120]}")

print("\n✓ All fixed!")

