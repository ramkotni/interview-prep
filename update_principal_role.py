from docx import Document
import os

def update_role_and_objective(file_path, role_type):
    """Update role titles and objectives for senior level positioning"""
    try:
        doc = Document(file_path)

        # Define objectives based on role type
        objectives = {
            'java': "Seeking a Principal Full Stack Engineer position to leverage 18 years of deep expertise in architecting and delivering enterprise-scale applications using Java, Spring Boot, microservices, cloud-native technologies, and modern frontend frameworks. Proven track record of technical leadership, driving architectural decisions, and mentoring teams to deliver high-quality, scalable solutions.",

            'python': "Seeking a Principal Full Stack Engineer position to leverage 18 years of deep expertise in architecting and delivering enterprise applications using Python, Django, FastAPI, React/Angular, and cloud technologies. Demonstrated capability in technical leadership, architectural innovation, and cross-functional team execution.",

            'agile': "Seeking a Principal Software Architect / Technical Lead position to leverage 18 years of deep expertise in enterprise PLM systems, complex integrations, and digital transformation programs. Proven ability to design robust architectures, lead technical initiatives, and drive enterprise-scale implementations."
        }

        # Determine which objective to use
        if 'Python' in file_path:
            new_objective = objectives['python']
        elif 'Agile' in file_path or 'PLM' in file_path:
            new_objective = objectives['agile']
        else:
            new_objective = objectives['java']

        # Update objectives
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == 'Objective:':
                if i+1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i+1]

                    # Clear and set new objective
                    for run in next_para.runs:
                        run._element.getparent().remove(run._element)

                    next_para.add_run(new_objective)
                    break

        doc.save(file_path)
        print(f"✓ Updated: {os.path.basename(file_path)}")
        return True

    except Exception as e:
        print(f"✗ Error: {os.path.basename(file_path)} - {str(e)}")
        return False

# All resume files
files = [
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram-Kotni-AgilePLM.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram-Kotni-AgilePLM.docx"
]

print("\nUPDATING TO PRINCIPAL/ARCHITECT LEVEL:\n")
print("="*80)

for file_path in files:
    if os.path.exists(file_path):
        if 'Python' in file_path:
            update_role_and_objective(file_path, 'python')
        elif 'Agile' in file_path or 'PLM' in file_path:
            update_role_and_objective(file_path, 'agile')
        else:
            update_role_and_objective(file_path, 'java')

print("="*80)
print("\n✓ All resumes updated to Principal/Architect level positioning!")
print("\nNEW POSITIONING:")
print("  • Java: Principal Full Stack Engineer")
print("  • Python: Principal Full Stack Engineer")
print("  • Agile PLM: Principal Software Architect / Technical Lead")

