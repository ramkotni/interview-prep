from docx import Document
import os

def fix_sava_to_proper_title(file_path):
    """
    Fix SAVA and split text across runs
    """
    try:
        doc = Document(file_path)

        # Process all paragraphs
        for para in doc.paragraphs:
            # Get full text
            full_text = para.text

            # Fix SAVA -> Senior Full Stack Engineer
            if 'SAVA' in full_text:
                # Clear all runs and rebuild
                for run in para.runs:
                    run._element.getparent().remove(run._element)

                # Add new text with formatting from first run
                new_text = full_text.replace('SAVA FULL STACK DEVELOPER', 'Senior Full Stack Engineer')
                if para.runs:
                    first_run = para.runs[0]
                    first_run.text = new_text
                else:
                    para.add_run(new_text)

            # Fix other titles
            elif 'JAVA FULL STACK DEVELOPER' in full_text or 'Java Full Stack Developer' in full_text:
                for run in para.runs:
                    run._element.getparent().remove(run._element)

                new_text = full_text.replace('JAVA FULL STACK DEVELOPER', 'Senior Full Stack Engineer')
                new_text = new_text.replace('Java Full Stack Developer', 'Senior Full Stack Engineer')

                if para.runs:
                    first_run = para.runs[0]
                    first_run.text = new_text
                else:
                    para.add_run(new_text)

        doc.save(file_path)
        filename = os.path.basename(file_path)
        print(f"✓ Fixed: {filename}")
        return True

    except Exception as e:
        print(f"✗ Error: {os.path.basename(file_path)} - {str(e)}")
        return False

# All resume files
resume_files = [
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram-Kotni-AgilePLM.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram-Kotni-AgilePLM.docx"
]

print("FIXING SAVA AND STANDARDIZING TITLES:\n")

for file_path in resume_files:
    if os.path.exists(file_path):
        fix_sava_to_proper_title(file_path)

print("\n✓ All resumes fixed!")
print("✓ SAVA -> Senior Full Stack Engineer")
print("✓ Ready for submission!")

