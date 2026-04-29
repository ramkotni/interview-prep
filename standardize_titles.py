from docx import Document
import os

def standardize_title(file_path):
    """Standardize all resume titles to 'Senior Full Stack Engineer'"""
    try:
        doc = Document(file_path)

        # Scan for title paragraph (usually para 1)
        for i, para in enumerate(doc.paragraphs[:10]):
            text = para.text.strip()

            # Check if this is a title line (contains developer/engineer/architect keywords)
            if any(x in text.upper() for x in ['DEVELOPER', 'ENGINEER', 'ARCHITECT', 'CONSULTANT']):
                # Only update if it's not already correct
                if 'SENIOR FULL STACK ENGINEER' not in text.upper():
                    # Clear all runs and set new text
                    for run in para.runs:
                        run._element.getparent().remove(run._element)

                    # Add single run with correct title
                    para.add_run('Senior Full Stack Engineer')
                    print(f"✓ Updated: {os.path.basename(file_path)}")
                    doc.save(file_path)
                    return True

        print(f"✓ Already correct: {os.path.basename(file_path)}")
        return True

    except Exception as e:
        print(f"✗ Error: {os.path.basename(file_path)} - {str(e)}")
        return False

# All files to standardize
files = [
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram-Kotni-AgilePLM.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram-Kotni-AgilePLM.docx"
]

print("Standardizing all resume titles to 'Senior Full Stack Engineer'...\n")

for file_path in files:
    if os.path.exists(file_path):
        standardize_title(file_path)

print("\n✓ All titles standardized!")

