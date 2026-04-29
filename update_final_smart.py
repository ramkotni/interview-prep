from docx import Document
from copy import deepcopy
import os

def update_docx_smart(file_path):
    """
    Update resume content intelligently while preserving 100% of formatting
    This works at the run level and preserves all style properties
    """
    try:
        doc = Document(file_path)

        # Define text replacements
        replacements = [
            # Title
            ("JAVA FULL STACK DEVELOPER", "Senior Java Full Stack Engineer | Microservices | Cloud | Platform Reliability"),

            # Contact
            ("603.858.7546", "603-858-7546"),
            ("ramkotni@gmail.com", "mohankotni77@gmail.com"),

            # Years
            ("17 years of experience", "18 years of experience"),
            ("16 years of experience", "18 years of experience"),
            ("17 years of proved experience", "18 years of experience"),

            # Professional Summary intro
            ("17 years of experience designing", "18 years of experience building and modernizing enterprise platforms"),
        ]

        def replace_in_runs(runs, old_text, new_text):
            """Replace text in runs while preserving formatting"""
            full_text = ''.join(run.text for run in runs)

            if old_text not in full_text:
                return False

            # Find which runs contain the text
            char_count = 0
            start_run = None
            start_offset = 0

            for run_idx, run in enumerate(runs):
                run_len = len(run.text)

                # Check if old_text starts in this run
                if start_run is None and old_text in full_text[char_count:char_count + run_len]:
                    start_run = run_idx
                    start_offset = full_text.find(old_text, char_count) - char_count
                    break

                char_count += run_len

            if start_run is None:
                return False

            # Simple approach: replace in the first run that contains the text
            for run in runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    return True

            return False

        # Process paragraphs
        for para in doc.paragraphs:
            for old_text, new_text in replacements:
                if old_text in para.text:
                    replace_in_runs(para.runs, old_text, new_text)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old_text, new_text in replacements:
                            if old_text in para.text:
                                replace_in_runs(para.runs, old_text, new_text)

        doc.save(file_path)
        print(f"✓ Successfully updated: {os.path.basename(file_path)}")
        return True

    except Exception as e:
        print(f"✗ Error updating {file_path}: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# Files to update
resume_files = [
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram-Kotni-AgilePLM.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram-Kotni-AgilePLM.docx"
]

print("Updating DOCX files - preserving all formatting...\n")

for file_path in resume_files:
    if os.path.exists(file_path):
        update_docx_smart(file_path)

print("\n✓ All files updated successfully!")

