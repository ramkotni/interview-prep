from docx import Document
from copy import deepcopy

TARGET_FILES = [
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\FINAL-RESUMES-2026\Ram-Kotni-Java-FullStack-Lead.docx",
]


def set_text_preserve_run_style(paragraph, text):
    first_run = paragraph.runs[0] if paragraph.runs else None
    style_props = None
    if first_run:
        f = first_run.font
        style_props = {
            "name": f.name,
            "size": f.size,
            "bold": f.bold,
            "italic": f.italic,
            "underline": f.underline,
            "color": f.color.rgb if f.color and f.color.rgb else None,
        }

    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    new_run = paragraph.add_run(text)
    if style_props:
        new_run.font.name = style_props["name"]
        new_run.font.size = style_props["size"]
        new_run.font.bold = style_props["bold"]
        new_run.font.italic = style_props["italic"]
        new_run.font.underline = style_props["underline"]
        if style_props["color"]:
            new_run.font.color.rgb = style_props["color"]


def first_index(paragraphs, predicate):
    for i, p in enumerate(paragraphs):
        if predicate((p.text or "").strip()):
            return i
    return None


def update_file(path):
    doc = Document(path)
    paras = doc.paragraphs

    # Header/title/contact cleanup
    title_idx = first_index(paras, lambda t: "full stack" in t.lower() and i_contains_any(t, ["engineer", "developer"]))
    if title_idx is not None:
        set_text_preserve_run_style(paras[title_idx], "Senior Java Full Stack Engineer")

    contact_idx = first_index(paras, lambda t: "@" in t and "linkedin.com" in t)
    if contact_idx is not None:
        set_text_preserve_run_style(
            paras[contact_idx],
            "Austin, TX  603-858-7546  mohankotni77@gmail.com  https://www.linkedin.com/in/ramkotni/",
        )

    # Objective -> Executive Summary
    obj_idx = first_index(paras, lambda t: t.lower().startswith("objective"))
    if obj_idx is not None:
        set_text_preserve_run_style(paras[obj_idx], "Executive Summary:")
        if obj_idx + 1 < len(paras):
            set_text_preserve_run_style(
                paras[obj_idx + 1],
                "Senior Java Full Stack Engineer with 18 years of experience building scalable, cloud-native, and high-performance enterprise systems. Expert in Java 17+, Spring Boot, JAX-RS, Microservices, Angular, Kafka, and distributed system design. Proven track record delivering mission-critical applications across energy, robotics, and healthcare domains. Strong in backend engineering, API architecture, AWS/GCP deployments, and DevOps automation, with consistent delivery of reliability and performance improvements.",
            )

    # Top summary/competencies section refresh
    ps_idx = first_index(paras, lambda t: t.lower().startswith("professional summary"))
    if ps_idx is not None:
        set_text_preserve_run_style(paras[ps_idx], "Core Competencies")

    # Replace stale summary line with ATS keyword block
    if ps_idx is not None and ps_idx + 2 < len(paras):
        set_text_preserve_run_style(
            paras[ps_idx + 2],
            "Java 17+ • Spring Boot • JAX-RS • Microservices • Angular 16+ • Kafka • REST API Architecture • AWS/GCP • Kubernetes • Docker • CI/CD (Jenkins, GitHub Actions) • Oracle/PostgreSQL • Redis • Event-Driven Systems • Performance Optimization • System Design • Technical Leadership",
        )

    core_idx = first_index(paras, lambda t: t.lower().startswith("core strengths"))
    if core_idx is not None:
        set_text_preserve_run_style(paras[core_idx], "Professional Summary")
        summary_lines = [
            "Architect-level engineer with a strong backend focus (70%) on Java, Spring Boot, microservices, secure API design, and performance tuning.",
            "Frontend capability (30%) includes Angular 16+, TypeScript, and responsive UI engineering for enterprise workflows.",
            "Hands-on cloud and DevOps expertise across AWS/GCP, Kubernetes, Docker, and CI/CD automation.",
            "Experienced in event-driven architecture using Kafka with strong observability practices (Prometheus, Kibana, Splunk).",
            "Proven technical leader in modernization programs, architecture reviews, mentoring, and Agile delivery.",
        ]
        for i, line in enumerate(summary_lines, start=1):
            if core_idx + i < len(paras):
                set_text_preserve_run_style(paras[core_idx + i], line)

    # Fix noisy role titles/typos
    ercot_role_idx = first_index(paras, lambda t: "apr 2025" in t.lower() and "ercot" not in t.lower())
    if ercot_role_idx is not None:
        set_text_preserve_run_style(paras[ercot_role_idx], "Sr. Java Full Stack Developer  |  Apr 2025 - Present")

    amz_role_idx = first_index(paras, lambda t: "feb 2023" in t.lower() and ("sava" in t.lower() or "steck" in t.lower() or "full" in t.lower()))
    if amz_role_idx is not None:
        set_text_preserve_run_style(paras[amz_role_idx], "Sr. Java Full Stack Developer  |  Feb 2023 - Mar 2025")

    # ERCOT bullet tightening (6-7 recruiter-friendly bullets)
    start_idx = first_index(paras, lambda t: t.startswith("Designed and implemented RESTful APIs with JAX-RS"))
    end_idx = first_index(paras, lambda t: t.startswith("Collaborated in design reviews and defect triage"))
    if start_idx is not None and end_idx is not None and end_idx >= start_idx:
        new_bullets = [
            "Designed and optimized JAX-RS APIs for grid operations, reducing processing latency and improving reliability.",
            "Enhanced Angular UI for RIOO-IS/RS across 10+ modules, improving navigation and user experience.",
            "Delivered a production-ready INR email notification system used across ERCOT operations.",
            "Refactored shared components for 10+ entities, reducing code duplication by 30%.",
            "Automated deployments using Ansible, enabling zero-downtime releases in higher environments.",
            "Resolved 25+ production issues and 50+ data corrections, improving SLA compliance and data accuracy.",
            "Built unit and integration tests for 20+ APIs, supporting 99% service reliability.",
        ]
        slot_count = end_idx - start_idx + 1
        for i in range(slot_count):
            text = new_bullets[i] if i < len(new_bullets) else ""
            set_text_preserve_run_style(paras[start_idx + i], text)

    doc.save(path)
    print(f"updated: {path}")


def i_contains_any(text, words):
    low = text.lower()
    return any(w.lower() in low for w in words)


if __name__ == "__main__":
    for p in TARGET_FILES:
        update_file(p)

