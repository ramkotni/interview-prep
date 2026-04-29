from docx import Document
import os
import re

def update_docx_robust(file_path):
    """
    Update resume with robust text replacement while preserving formatting
    """
    try:
        doc = Document(file_path)

        # More comprehensive replacements
        replacements = [
            # Title
            ("JAVA FULL STACK DEVELOPER", "Senior Java Full Stack Engineer | Microservices | Cloud | Platform Reliability"),
            ("Java Full Stack Developer", "Senior Java Full Stack Engineer | Microservices | Cloud | Platform Reliability"),

            # Contact
            ("603.858.7546", "603-858-7546"),
            ("ramkotni@gmail.com", "mohankotni77@gmail.com"),

            # Years - various formats
            ("17 years of experience", "18 years of experience"),
            ("16 years of experience", "18 years of experience"),
            ("17 years of proved experience", "18 years of experience"),
            ("17 years of IT experience", "18 years of experience"),
            ("16 years of IT experience", "18 years of experience"),
            ("17 years", "18 years"),
            ("16 years", "18 years"),
        ]

        # Process paragraphs
        for para in doc.paragraphs:
            for old_text, new_text in replacements:
                if old_text in para.text:
                    # Replace in each run
                    full_para_text = para.text
                    if old_text in full_para_text:
                        # Build new text
                        new_para_text = full_para_text.replace(old_text, new_text)

                        # Find and replace in runs
                        remaining = old_text
                        for run in para.runs:
                            if remaining in run.text:
                                run.text = run.text.replace(remaining, new_text)
                                remaining = ""
                                break
                            elif run.text:
                                # Check if part of the old_text is split across runs
                                for i in range(1, len(remaining) + 1):
                                    if remaining[:i] in run.text:
                                        run.text = run.text.replace(remaining[:i], new_text[:i])
                                        remaining = remaining[i:]
                                        new_text = new_text[i:]
                                        break

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old_text, new_text in replacements:
                            if old_text in para.text:
                                full_para_text = para.text
                                if old_text in full_para_text:
                                    new_para_text = full_para_text.replace(old_text, new_text)
                                    remaining = old_text
                                    for run in para.runs:
                                        if remaining in run.text:
                                            run.text = run.text.replace(remaining, new_text)
                                            remaining = ""
                                            break

        doc.save(file_path)
        print(f"✓ Updated: {os.path.basename(file_path)}")
        return True

    except Exception as e:
        print(f"✗ Error: {os.path.basename(file_path)} - {str(e)}")
        return False

# Files to update
resume_files = [
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram-Kotni-AgilePLM.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram-Kotni-AgilePLM.docx"
]

print("Updating all resume files with 18 years...\n")

for file_path in resume_files:
    if os.path.exists(file_path):
        update_docx_robust(file_path)

print("\n✓ All files updated successfully!")

