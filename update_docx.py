from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

# ATS-friendly resume content
resume_content = {
    "name": "Ram Mohan Kotni",
    "location": "Austin, TX | 603-858-7546 | mohankotni77@gmail.com | https://www.linkedin.com/in/ramkotni/",
    "title": "Senior Java Full Stack Engineer | Microservices | Cloud | Platform Reliability",
    "summary": "Senior Java Full Stack Engineer with 18 years of experience building and modernizing enterprise platforms across Java, Spring Boot, microservices, Angular, and cloud-native environments. Strong track record in production reliability, secure API design, event-driven architecture, and CI/CD automation. Hands-on leader known for resolving critical incidents, improving performance, and delivering measurable business outcomes in regulated and high-throughput domains.",
    "skills": {
        "Backend": "Java 8/11/17, Spring Boot, Spring Security, REST APIs, Microservices, JPA/Hibernate",
        "Frontend": "Angular, React, TypeScript, JavaScript, HTML5, CSS3",
        "Cloud and DevOps": "AWS (EC2, S3, Lambda, RDS, CloudWatch, IAM), GCP (GKE, Cloud Run), Docker, Kubernetes, Jenkins, GitLab CI/CD, Ansible",
        "Data and Messaging": "Oracle, PostgreSQL, MySQL, MongoDB, Cassandra, Redis, Kafka",
        "Testing and Quality": "JUnit, Mockito, Postman, Swagger/OpenAPI, SonarQube",
        "Observability": "CloudWatch, Prometheus, Elasticsearch, Kibana, Splunk",
        "Ways of Working": "Agile Scrum, architecture reviews, production support, mentoring"
    },
    "experience": [
        {
            "company": "ERCOT (Electric Reliability Council of Texas), Taylor, TX",
            "title": "Sr. Java Full Stack Developer",
            "dates": "Apr 2025 - Present",
            "bullets": [
                "Delivered and stabilized production workflows for resource submission and approval use cases in ERCOT systems.",
                "Refactored shared code across 10+ entities, reducing duplication and improving maintainability.",
                "Built 10+ backend APIs and integrated Angular routing updates for role-based user flows.",
                "Resolved 25+ production incidents within SLA and closed 50+ data correction tickets with high accuracy.",
                "Improved deployment reliability by using Ansible-driven deployment and validation checks across environments."
            ]
        },
        {
            "company": "Amazon Robotics, Boston, US",
            "title": "Sr. Java Full Stack Developer",
            "dates": "Feb 2023 - Mar 2025",
            "bullets": [
                "Designed and delivered Spring Boot microservices for manufacturing, tracking, and warehouse operations.",
                "Built Angular dashboards for real-time operational visibility, reducing manual tracking effort.",
                "Implemented Kafka-based event-driven flows, improving service-to-service communication latency.",
                "Improved system availability and release speed using AWS services, CI/CD automation, and observability.",
                "Strengthened security with OAuth2/JWT patterns and improved testing depth with unit and integration coverage."
            ]
        },
        {
            "company": "Biogen, North Carolina, US",
            "title": "Lead Full Stack Java Developer",
            "dates": "Jun 2022 - Jan 2023",
            "bullets": [
                "Built secure, compliant full-stack workflows for supply and compliance data handling.",
                "Developed Spring Boot APIs and Angular dashboards for real-time status and analytics.",
                "Integrated internal and external systems to improve traceability and operational efficiency.",
                "Contributed to cloud optimization and improved uptime through monitoring and performance tuning."
            ]
        },
        {
            "company": "Dell Technologies",
            "title": "Application Architect / Technical Lead",
            "dates": "Jul 2015 - May 2022",
            "bullets": [
                "Led migration from monolith to Spring Boot microservices and improved deployment frequency.",
                "Architected enterprise UI and API layers for scalable product lifecycle and integration workflows.",
                "Drove CI/CD adoption, code quality practices, and design standards across teams.",
                "Mentored developers and led architecture and code reviews for critical releases."
            ]
        },
        {
            "company": "IBM, New York, US",
            "title": "Sr. Java Developer",
            "dates": "May 2012 - Jul 2015",
            "bullets": [
                "Developed and enhanced enterprise Java modules for packaging and barcode workflows.",
                "Improved UI, API integration, and code quality practices across releases.",
                "Delivered features in Agile teams while improving quality and maintainability."
            ]
        },
        {
            "company": "Innominds, San Jose, CA",
            "title": "Java Developer",
            "dates": "Dec 2009 - Mar 2012",
            "bullets": [
                "Built Java and PL/SQL solutions for workflow, integration, and data processing use cases.",
                "Delivered Spring/SOA services, performance testing scripts, and quality improvements."
            ]
        },
        {
            "company": "Wells Fargo, IA",
            "title": "Java Developer",
            "dates": "Nov 2007 - May 2009",
            "bullets": [
                "Built Java and PL/SQL solutions for workflow, integration, and data processing use cases.",
                "Delivered Spring/SOA services, performance testing scripts, and quality improvements."
            ]
        }
    ],
    "certifications": [
        "AWS Certified Solutions Architect - Associate",
        "Oracle Certified Professional, Java SE Programmer",
        "Agile Scrum Master Certification"
    ],
    "education": [
        "Master of Computer Applications from MK University, 2001",
        "Bachelor of Science (Mathematics) from Andhra University, 1998"
    ],
    "achievements": [
        "Improved reliability through production monitoring and performance optimization practices.",
        "Increased deployment velocity through CI/CD and release automation.",
        "Improved user and stakeholder experience through better workflow visibility and response times.",
        "Resolved 25+ production incidents within SLA at ERCOT with high accuracy."
    ]
}

def clear_paragraph(paragraph):
    """Clear a paragraph while preserving formatting"""
    for run in paragraph.runs:
        run.text = ""

def update_docx_file(file_path):
    """Update a docx file with ATS-friendly content while preserving formatting"""
    try:
        doc = Document(file_path)

        # Keep only the first 5 paragraphs for header/title info, then clear the rest
        for para in doc.paragraphs[5:]:
            clear_paragraph(para)

        # Update content - find sections and replace them
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # Update name
            if "Ram" in text and "Kotni" in text and i < 5:
                para.clear()
                run = para.add_run("Ram Mohan Kotni")

            # Update location/contact
            elif "Austin" in text or "603" in text:
                para.clear()
                run = para.add_run(resume_content["location"])

            # Update title/headline
            elif "JAVA FULL STACK" in text or "Java Full Stack" in text:
                para.clear()
                run = para.add_run(resume_content["title"])

            # Update Professional Summary
            elif "Professional Summary" in text or "Objective" in text:
                para.clear()
                run = para.add_run("Professional Summary")

                # Add next paragraph with summary
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    next_para.clear()
                    next_para.add_run(resume_content["summary"])

        doc.save(file_path)
        print(f"✓ Updated: {file_path}")
        return True
    except Exception as e:
        print(f"✗ Error updating {file_path}: {str(e)}")
        return False

# List of resume files to update
resume_files = [
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram-Kotni-AgilePLM.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram-Kotni-AgilePLM.docx"
]

print("Starting to update DOCX files...\n")
for file_path in resume_files:
    if os.path.exists(file_path):
        update_docx_file(file_path)
    else:
        print(f"✗ File not found: {file_path}")

print("\nUpdate complete!")

