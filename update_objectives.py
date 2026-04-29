from docx import Document

file = r'C:\RamKotni\Personal\interview-prep\resumework\Ram Mohan Kotni-Java-FullStack Lead.docx'
doc = Document(file)

# Find the Objective section and update it
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == 'Objective:':
        # The next paragraph should contain the objective text
        if i+1 < len(doc.paragraphs):
            next_para = doc.paragraphs[i+1]

            # New objective based on cover letter and role
            new_objective = "Seeking a Java Full Stack Developer position to leverage 18 years of experience in designing and developing scalable enterprise applications using Java, Spring Boot, microservices, cloud technologies, and modern frontend frameworks. Committed to delivering high-quality solutions through agile methodologies, clean code practices, and cross-functional collaboration."

            # Clear runs and set new text
            for run in next_para.runs:
                run._element.getparent().remove(run._element)

            next_para.add_run(new_objective)
            print(f"Updated objective at paragraph {i+1}")
            break

# Also update both Resume folder copies
files_to_update = [
    r'C:\RamKotni\Personal\interview-prep\Resume\Ram Mohan Kotni-Java-FullStack Lead.docx',
    r'C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx',
    r'C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx',
    r'C:\RamKotni\Personal\interview-prep\resumework\Ram-Kotni-AgilePLM.docx',
    r'C:\RamKotni\Personal\interview-prep\Resume\Ram-Kotni-AgilePLM.docx'
]

for file_path in files_to_update:
    try:
        doc = Document(file_path)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == 'Objective:':
                if i+1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i+1]

                    # Different objectives based on file type
                    if 'Python' in file_path:
                        new_objective = "Seeking a Python Full Stack Developer position to leverage 18 years of experience in designing and developing scalable applications using Python, Django, FastAPI, React/Angular, cloud technologies. Committed to delivering high-quality solutions through agile methodologies and cross-functional collaboration."
                    elif 'Agile' in file_path or 'PLM' in file_path:
                        new_objective = "Seeking an Agile PLM / Enterprise Integration role to leverage 18 years of experience in PLM systems, enterprise architecture, Java-based integrations, and modernization programs. Committed to delivering robust solutions through agile practices and cross-functional collaboration."
                    else:
                        new_objective = "Seeking a Java Full Stack Developer position to leverage 18 years of experience in designing and developing scalable enterprise applications using Java, Spring Boot, microservices, cloud technologies, and modern frontend frameworks. Committed to delivering high-quality solutions through agile methodologies, clean code practices, and cross-functional collaboration."

                    for run in next_para.runs:
                        run._element.getparent().remove(run._element)

                    next_para.add_run(new_objective)
                    print(f"Updated: {file_path.split(chr(92))[-1]}")
                    break

        doc.save(file_path)
    except Exception as e:
        print(f"Error with {file_path}: {e}")

# Save the first file
doc = Document(file)
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == 'Objective:':
        if i+1 < len(doc.paragraphs):
            next_para = doc.paragraphs[i+1]
            new_objective = "Seeking a Java Full Stack Developer position to leverage 18 years of experience in designing and developing scalable enterprise applications using Java, Spring Boot, microservices, cloud technologies, and modern frontend frameworks. Committed to delivering high-quality solutions through agile methodologies, clean code practices, and cross-functional collaboration."

            for run in next_para.runs:
                run._element.getparent().remove(run._element)

            next_para.add_run(new_objective)
            break

doc.save(file)
print("\n✓ All objectives updated with cover letter alignment!")

