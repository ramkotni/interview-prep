from docx import Document
from docx.oxml import parse_xml
from copy import deepcopy
import os
import shutil

def backup_file(file_path):
    """Create a backup of the file"""
    backup_path = file_path.replace('.docx', '.backup_ats.docx')
    shutil.copy2(file_path, backup_path)
    return backup_path

def replace_text_in_runs(paragraph, old_text, new_text):
    """Replace text in paragraph while preserving formatting"""
    full_text = ''.join(run.text for run in paragraph.runs)

    if old_text in full_text:
        # Clear all runs
        for run in paragraph.runs:
            run.text = ""

        # Add new text to first run
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
        return True
    return False

def update_docx_comprehensive(file_path):
    """Comprehensively update docx file with new resume content"""
    try:
        doc = Document(file_path)

        # Define replacements - look for sections and update them
        replacements = {
            "JAVA FULL STACK DEVELOPER": "Senior Java Full Stack Engineer | Microservices | Cloud | Platform Reliability",
            "Austin, TX | 603.858.7546": "Austin, TX | 603-858-7546 | mohankotni77@gmail.com | https://www.linkedin.com/in/ramkotni/",
            "603.858.7546": "603-858-7546",
            "https://www.linkedin.com/in/mohankotni/": "https://www.linkedin.com/in/ramkotni/",
        }

        for para in doc.paragraphs:
            for old, new in replacements.items():
                replace_text_in_runs(para, old, new)

        # Also check in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for old, new in replacements.items():
                            replace_text_in_runs(para, old, new)

        doc.save(file_path)
        print(f"✓ Updated: {file_path}")
        return True
    except Exception as e:
        print(f"✗ Error updating {file_path}: {str(e)}")
        return False

# Main resume files to update
resume_files = [
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram-Kotni-AgilePLM.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram-Kotni-AgilePLM.docx"
]

print("Starting to update DOCX files with ATS-friendly content...\n")
updated_count = 0

for file_path in resume_files:
    if os.path.exists(file_path):
        backup_file(file_path)
        if update_docx_comprehensive(file_path):
            updated_count += 1
    else:
        print(f"✗ File not found: {file_path}")

print(f"\n✓ Successfully updated {updated_count} files!")
print("Backup files created with .backup_ats.docx extension")

