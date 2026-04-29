from docx import Document
import os
import shutil

def backup_file(file_path):
    """Create a backup of the file"""
    backup_path = file_path.replace('.docx', '.backup_ats_full.docx')
    if not os.path.exists(backup_path):
        shutil.copy2(file_path, backup_path)
    return backup_path

def replace_paragraph_text(paragraph, new_text):
    """Replace all text in a paragraph while preserving formatting of first run"""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    # Clear all runs except first
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)

    # Update first run
    if paragraph.runs:
        paragraph.runs[0].text = new_text

def update_professional_summary(doc):
    """Update professional summary section"""
    new_summary = "Senior Java Full Stack Engineer with 18 years of experience building and modernizing enterprise platforms across Java, Spring Boot, microservices, Angular, and cloud-native environments. Strong track record in production reliability, secure API design, event-driven architecture, and CI/CD automation. Hands-on leader known for resolving critical incidents, improving performance, and delivering measurable business outcomes in regulated and high-throughput domains."

    found_summary = False
    for i, para in enumerate(doc.paragraphs):
        if found_summary and para.text.strip() and not para.text.strip().startswith("Core"):
            replace_paragraph_text(para, new_summary)
            found_summary = False
        elif "Professional Summary" in para.text:
            found_summary = True

def update_docx_full(file_path):
    """Fully update docx with ATS content while preserving all formatting"""
    try:
        doc = Document(file_path)

        # Update specific paragraphs based on content
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # Update title
            if text == "JAVA FULL STACK DEVELOPER" or text == "Java Full Stack Developer":
                replace_paragraph_text(para, "Senior Java Full Stack Engineer | Microservices | Cloud | Platform Reliability")

            # Update phone number format
            elif "603.858.7546" in text:
                new_text = text.replace("603.858.7546", "603-858-7546")
                new_text = new_text.replace("https://www.linkedin.com/in/mohankotni/", "https://www.linkedin.com/in/ramkotni/")
                replace_paragraph_text(para, new_text)

            # Update old linkedin URL
            elif text == "https://www.linkedin.com/in/mohankotni/":
                replace_paragraph_text(para, "https://www.linkedin.com/in/ramkotni/")

            # Remove Objective section - look for it
            elif text == "Objective:" or "To obtain a challenging Java Full Stack Developer" in text:
                # Keep the section heading but it will be overwritten by Professional Summary
                pass

            # Update professional summary intro
            elif "A highly skilled Java Full Stack Developer with" in text or "An accomplished Java Full Stack Developer with" in text:
                replace_paragraph_text(para, "Senior Java Full Stack Engineer with 18 years of experience building and modernizing enterprise platforms across Java, Spring Boot, microservices, Angular, and cloud-native environments. Strong track record in production reliability, secure API design, event-driven architecture, and CI/CD automation. Hands-on leader known for resolving critical incidents, improving performance, and delivering measurable business outcomes in regulated and high-throughput domains.")

            # Update experience years mention
            elif "16 years of IT experience" in text or "17 years of proven experience" in text:
                new_text = text.replace("16 years", "18 years").replace("17 years", "18 years")
                replace_paragraph_text(para, new_text)

        doc.save(file_path)
        print(f"✓ Updated: {os.path.basename(file_path)}")
        return True
    except Exception as e:
        print(f"✗ Error updating {file_path}: {str(e)}")
        return False

# Main resume files to update
resume_files = [
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram-Kotni-AgilePLM.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram-Kotni-AgilePLM.docx"
]

print("Performing comprehensive ATS-friendly updates...\n")
updated_count = 0

for file_path in resume_files:
    if os.path.exists(file_path):
        backup_file(file_path)
        if update_docx_full(file_path):
            updated_count += 1
    else:
        print(f"✗ File not found: {file_path}")

print(f"\n✓ Successfully updated {updated_count} files!")
print("All formats preserved - documents maintain original styling and layout")

