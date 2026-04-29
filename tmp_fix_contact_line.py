from docx import Document

FILES = [
    r"C:\RamKotni\Personal\interview-prep\resumework\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\Resume\Ram Mohan Kotni-Java-FullStack Lead.docx",
    r"C:\RamKotni\Personal\interview-prep\FINAL-RESUMES-2026\Ram-Kotni-Java-FullStack-Lead.docx",
]

CONTACT = "Austin, TX  603-858-7546  mohankotni77@gmail.com  https://www.linkedin.com/in/ramkotni/"


def replace_paragraph_full(paragraph, text):
    p = paragraph._p
    for child in list(p):
        # preserve paragraph properties only
        if child.tag.endswith('}pPr'):
            continue
        p.remove(child)
    paragraph.add_run(text)


for path in FILES:
    doc = Document(path)
    for para in doc.paragraphs[:8]:
        t = (para.text or "").strip().lower()
        if "linkedin.com" in t or "@" in t or "603" in t:
            replace_paragraph_full(para, CONTACT)
            break
    doc.save(path)
    print(f"fixed contact: {path}")

