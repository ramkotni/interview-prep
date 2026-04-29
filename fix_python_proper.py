from docx import Document

files = [
    r'C:\RamKotni\Personal\interview-prep\resumework\Mohan Kotni-Python-FullStack Lead.docx',
    r'C:\RamKotni\Personal\interview-prep\Resume\Mohan Kotni-Python-FullStack Lead.docx'
]

for file in files:
    doc = Document(file)

    # Fix para 13 properly - rebuild it
    para = doc.paragraphs[13]

    # Get the full text and fix it
    full_text = para.text
    corrected_text = full_text.replace('13+', '18+').replace('over 1', 'over ').replace('over 18', 'over 18')

    # Clear all runs
    for run in para.runs:
        run._element.getparent().remove(run._element)

    # Add corrected text
    para.add_run(corrected_text)

    print(f"Fixed: {file.split(chr(92))[-1]}")
    print(f"  New text: {para.text[:80]}...")

    doc.save(file)

print("\n✓ Python resume years fixed!")

